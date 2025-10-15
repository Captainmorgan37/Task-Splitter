import re
from dataclasses import dataclass, field
from datetime import datetime, timedelta, date, time, timezone
from functools import lru_cache
from io import BytesIO
from pathlib import Path
from collections.abc import Mapping
from typing import List, Dict, Any, Tuple, Optional, Set, Sequence

import pandas as pd
import pytz
from zoneinfo import ZoneInfo
import streamlit as st
from pandas.api.types import is_scalar

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from fl3xx_api import (
    DEFAULT_FL3XX_BASE_URL,
    Fl3xxApiConfig,
    enrich_flights_with_crew,
    fetch_flights,
)

# ----------------------------
# App Config
# ----------------------------
st.set_page_config(page_title="Night-Shift Tail Splitter", layout="wide")
st.title("🛫 Night-Shift Tail Splitter")

st.caption(
    "Assign next-day tails to on-duty shifts as evenly as possible, while keeping all legs of a tail together."
)

UTC = timezone.utc
LOCAL_TZ = ZoneInfo("America/Edmonton")
AIRPORT_TZ_FILENAME = "Airport TZ.txt"
DEPARTURE_WINDOW_START_UTC = time(hour=8, tzinfo=UTC)
DEPARTURE_WINDOW_END_UTC = time(hour=8, tzinfo=UTC)

# ----------------------------
# Types
# ----------------------------
@dataclass
class TailPackage:
    tail: str
    legs: int
    workload: float
    first_local_dt: datetime  # first dep local datetime for the day
    sample_legs: List[Dict[str, Any]]  # optional preview rows for UI (subset)
    has_priority: bool = False
    priority_labels: List[str] = field(default_factory=list)
    customs_legs: int = 0


# ----------------------------
# Helpers
# ----------------------------
_CUSTOMS_WORKLOAD_MULTIPLIER = 1.5

_DEPARTURE_AIRPORT_COLUMNS: Sequence[str] = (
    "departure_airport",
    "dep_airport",
    "departureAirport",
    "departure_airport_code",
    "airportFrom",
    "fromAirport",
)

_ARRIVAL_AIRPORT_COLUMNS: Sequence[str] = (
    "arrival_airport",
    "arr_airport",
    "arrivalAirport",
    "arrival_airport_code",
    "airportTo",
    "toAirport",
)


def _safe_parse_dt(dt_str: str) -> datetime:
    """Parse ISO-like datetime. If timezone-naive, assume UTC."""
    try:
        dt = datetime.fromisoformat(dt_str.replace("Z", "+00:00"))
        if dt.tzinfo is None:
            return dt.replace(tzinfo=pytz.UTC)
        return dt
    except Exception:
        # Last resort: try pandas
        dt = pd.to_datetime(dt_str, utc=True).to_pydatetime()
        return dt


def _to_local(dt: datetime, tz_name: str | None) -> datetime:
    if tz_name:
        try:
            return dt.astimezone(ZoneInfo(tz_name))
        except Exception:
            pass
    # Fallback: leave in original tz; if naive, assume UTC then convert to LOCAL_TZ so ordering is at least consistent
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=pytz.UTC)
    return dt.astimezone(LOCAL_TZ)


def _priority_label(value: Any) -> Optional[str]:
    if value is None:
        return None
    if isinstance(value, str):
        text = value.strip()
    else:
        text = str(value).strip()
    if not text:
        return None
    if "priority" in text.lower():
        return text
    return None


def _normalize_person_name(value: Any) -> str:
    if value is None:
        return ""
    if not is_scalar(value):
        # Lists/dicts sometimes store crew metadata; ignore them here.
        return ""
    try:
        if pd.isna(value):
            return ""
    except (TypeError, ValueError):
        # Non-scalar objects or custom classes may raise here; treat them as empty.
        pass
    if not value:
        return ""
    text = str(value).strip()
    return text


def _member_display_name(member: Mapping[str, Any]) -> str:
    candidates = [
        member.get(key)
        for key in (
            "displayName",
            "display_name",
            "name",
            "fullName",
            "full_name",
        )
        if isinstance(member, Mapping)
    ]
    for candidate in candidates:
        name = _normalize_person_name(candidate)
        if name:
            return name
    if isinstance(member, Mapping):
        first = _normalize_person_name(
            member.get("firstName") or member.get("first_name")
        )
        last = _normalize_person_name(
            member.get("lastName") or member.get("last_name")
        )
        combined = " ".join(part for part in (first, last) if part)
        if combined.strip():
            return combined.strip()
    return ""


_SUBCHARTER_PATTERN = re.compile(r"subcharter", re.IGNORECASE)


def _workflow_indicates_subcharter(value: Any) -> bool:
    if value is None:
        return False
    return bool(_SUBCHARTER_PATTERN.search(str(value)))


def _filter_out_subcharter_rows(
    rows: List[Dict[str, Any]]
) -> Tuple[List[Dict[str, Any]], int]:
    filtered: List[Dict[str, Any]] = []
    skipped = 0
    for row in rows:
        workflow_value: Optional[Any] = None
        if isinstance(row, Mapping):
            for key in (
                "workflowCustomName",
                "workflow_custom_name",
                "workflowName",
                "workflow",
            ):
                if key in row and row[key] not in (None, ""):
                    workflow_value = row[key]
                    break
        if _workflow_indicates_subcharter(workflow_value):
            skipped += 1
            continue
        filtered.append(row)
    return filtered, skipped


_PIC_KEYWORDS = {
    "pic",
    "picname",
    "captain",
    "pilotincommand",
    "pilot_in_command",
    "pilotcommand",
}

_SIC_KEYWORDS = {
    "sic",
    "sicname",
    "copilot",
    "firstofficer",
    "first_officer",
}


def _crew_names_from_row(row: Mapping[str, Any]) -> Tuple[str, str]:
    pic = ""
    sic = ""
    for key, value in row.items():
        if value is None:
            continue
        normalized_key = re.sub(r"[^a-z]", "", str(key).lower())
        if not normalized_key:
            continue
        name = _normalize_person_name(value)
        if not name:
            continue
        if not pic and normalized_key in _PIC_KEYWORDS:
            pic = name
        elif not sic and normalized_key in _SIC_KEYWORDS:
            sic = name
    crew_members = row.get("crewMembers")
    if isinstance(crew_members, list):
        for member in crew_members:
            if not isinstance(member, Mapping):
                continue
            role = str(member.get("role") or member.get("position") or "").lower()
            is_pic = bool(member.get("isPIC") or "pic" in role)
            is_sic = bool(member.get("isSIC") or "sic" in role or "first officer" in role)
            name = _member_display_name(member)
            if name:
                if not pic and is_pic:
                    pic = name
                elif not sic and is_sic:
                    sic = name
    return pic, sic


def _crew_names_from_package(pkg: "TailPackage") -> Tuple[str, str]:
    pic = ""
    sic = ""
    for leg in pkg.sample_legs:
        if isinstance(leg, Mapping):
            leg_pic, leg_sic = _crew_names_from_row(leg)
            if not pic and leg_pic:
                pic = leg_pic
            if not sic and leg_sic:
                sic = leg_sic
            if pic and sic:
                break
    return pic, sic


_TAIL_PLACEHOLDER_PREFIXES = ("ADD", "NEW", "TBD", "TEMP", "HOLD", "UNKNOWN", "UNK")
_TAIL_PLACEHOLDER_VALUES = {"", "NA", "N/A", "NONE", "NULL", "-"}
_TAIL_US_PATTERN = re.compile(r"^N[0-9]{1,5}[A-Z]{0,2}$")
_TAIL_HYPHEN_PATTERN = re.compile(r"^[A-Z0-9]{1,2}-[A-Z0-9]{2,5}$")
_TAIL_ALNUM_PATTERN = re.compile(r"^[A-Z0-9]{4,7}$")


def _is_valid_tail_registration(value: Any) -> bool:
    if not isinstance(value, str):
        return False
    candidate = value.strip().upper()
    if not candidate or candidate in _TAIL_PLACEHOLDER_VALUES:
        return False
    if any(ch.isspace() for ch in candidate):
        return False
    if candidate.startswith(_TAIL_PLACEHOLDER_PREFIXES):
        return False
    if len(candidate) < 3:
        return False
    if _TAIL_US_PATTERN.fullmatch(candidate):
        return True
    if _TAIL_HYPHEN_PATTERN.fullmatch(candidate):
        return True
    if "-" not in candidate and not any(ch.isdigit() for ch in candidate):
        return False
    if _TAIL_ALNUM_PATTERN.fullmatch(candidate):
        return True
    return False


def _default_target_date() -> date:
    """Return the default target date (two days ahead in local Mountain time)."""
    now_local = datetime.now(LOCAL_TZ)
    return (now_local + timedelta(days=2)).date()


def _default_shift_labels(count: int) -> List[str]:
    presets: Dict[int, List[str]] = {
        3: ["0500", "0800", "1200"],
        4: ["0500", "0600", "0800", "1200"],
        5: ["0500", "0600", "0800", "0900", "1200"],
    }
    if count in presets:
        return list(presets[count])
    return [f"Shift {i+1}" for i in range(count)]


def _compute_departure_window_bounds(target_date: date) -> Tuple[datetime, datetime]:
    start = datetime.combine(target_date, DEPARTURE_WINDOW_START_UTC)
    end_date = target_date + timedelta(days=1)
    end = datetime.combine(end_date, DEPARTURE_WINDOW_END_UTC)
    return start, end


def _format_utc(dt: datetime) -> str:
    return dt.astimezone(UTC).isoformat().replace("+00:00", "Z")


def _filter_rows_by_departure_window(
    rows: List[Dict[str, Any]],
    start_utc: datetime,
    end_utc: datetime,
) -> Tuple[List[Dict[str, Any]], Dict[str, int]]:
    stats = {
        "total": len(rows),
        "within_window": 0,
        "before_window": 0,
        "after_window": 0,
    }
    if not rows:
        return [], stats

    filtered: List[Dict[str, Any]] = []

    for row in rows:
        dep_raw = row.get("dep_time")
        if dep_raw is None:
            stats["before_window"] += 1
            continue
        dt = _safe_parse_dt(str(dep_raw))
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=UTC)
        else:
            dt = dt.astimezone(UTC)
        if dt < start_utc:
            stats["before_window"] += 1
            continue
        if dt > end_utc:
            stats["after_window"] += 1
            continue
        filtered.append(row)
        stats["within_window"] += 1

    return filtered, stats


def _airport_tz_path() -> Path:
    return Path(__file__).with_name(AIRPORT_TZ_FILENAME)


@lru_cache(maxsize=1)
def _load_airport_tz_lookup() -> Dict[str, str]:
    metadata = _load_airport_metadata_lookup()
    tz_lookup: Dict[str, str] = {}
    for code, record in metadata.items():
        if isinstance(record, Mapping):
            tz_value = record.get("tz")
            if isinstance(tz_value, str) and tz_value:
                tz_lookup[code] = tz_value
    return tz_lookup


@lru_cache(maxsize=1)
def _load_airport_metadata_lookup() -> Dict[str, Dict[str, Optional[str]]]:
    path = _airport_tz_path()
    if not path.exists():
        return {}
    try:
        df = pd.read_csv(path)
    except Exception:
        return {}

    lookup: Dict[str, Dict[str, Optional[str]]] = {}
    for _, row in df.iterrows():
        tz: Optional[str] = None
        tz_value = row.get("tz")
        if isinstance(tz_value, str) and tz_value.strip():
            tz = tz_value.strip()

        country: Optional[str] = None
        country_value = row.get("country")
        if isinstance(country_value, str) and country_value.strip():
            country = country_value.strip()

        if tz is None and country is None:
            continue
        for key in ("icao", "iata", "lid"):
            code_value = row.get(key)
            if isinstance(code_value, str) and code_value.strip():
                lookup[code_value.strip().upper()] = {
                    "tz": tz,
                    "country": country,
                }
    return lookup


def _extract_codes(value: Any) -> List[str]:
    if not isinstance(value, str):
        return []
    cleaned = value.strip()
    if not cleaned:
        return []
    upper = cleaned.upper()
    if upper.replace(" ", "").isalnum() and len(upper.strip()) in {3, 4}:
        return [upper]
    return [token.upper() for token in re.findall(r"\b[A-Za-z0-9]{3,4}\b", upper)]


def _airport_country_from_row(
    row: Mapping[str, Any],
    columns: Sequence[str],
    lookup: Dict[str, Dict[str, Optional[str]]],
) -> Optional[str]:
    for column in columns:
        value = row.get(column)
        if value is None:
            continue
        if isinstance(value, float) and pd.isna(value):
            continue
        for code in _extract_codes(str(value)):
            record = lookup.get(code)
            if not record:
                continue
            country = record.get("country") if isinstance(record, Mapping) else None
            if isinstance(country, str) and country.strip():
                return country.strip()
    return None


def _is_customs_leg(row: Mapping[str, Any], lookup: Dict[str, Dict[str, Optional[str]]]) -> bool:
    if not lookup:
        return False
    dep_country = _airport_country_from_row(row, _DEPARTURE_AIRPORT_COLUMNS, lookup)
    arr_country = _airport_country_from_row(row, _ARRIVAL_AIRPORT_COLUMNS, lookup)
    if dep_country and arr_country:
        return dep_country != arr_country
    return False


def _apply_airport_timezones(df: pd.DataFrame) -> Tuple[pd.DataFrame, Set[str], bool]:
    if df.empty:
        return df, set(), False
    if "dep_tz" not in df.columns:
        df["dep_tz"] = None

    lookup = _load_airport_tz_lookup()
    lookup_used = bool(lookup)

    missing: Set[str] = set()

    def _needs_timezone(val: Any) -> bool:
        if val is None:
            return True
        if isinstance(val, float) and pd.isna(val):
            return True
        if isinstance(val, str) and not val.strip():
            return True
        return False

    for idx, row in df.iterrows():
        if not _needs_timezone(row.get("dep_tz")):
            continue
        airport_value: Optional[str] = None
        for col in _DEPARTURE_AIRPORT_COLUMNS:
            if col in df.columns and not pd.isna(row.get(col)):
                airport_value = str(row[col])
                if airport_value:
                    break
        if not airport_value:
            continue
        codes = _extract_codes(airport_value)
        tz_guess = None
        if lookup_used:
            tz_guess = next((lookup.get(code) for code in codes if code in lookup), None)
        if tz_guess:
            df.at[idx, "dep_tz"] = tz_guess
        else:
            missing.add(airport_value)

    return df, missing, lookup_used


# ----------------------------
# Data Fetch (stub or real)
# ----------------------------
@st.cache_data(show_spinner=False)
def fetch_next_day_legs(
    target_date: date,
    *,
    use_stub: bool,
    fl3xx_settings: Optional[Dict[str, Any]] = None,
    fetch_crew: bool = False,
) -> Tuple[pd.DataFrame, Dict[str, Any], Optional[Dict[str, Any]]]:
    """
    Return a DataFrame of legs for target_date with columns at least:
      tail (str), leg_id (str/int), dep_time (ISO str), dep_tz (IANA tz name)
    You can extend with more columns if your API provides (dep_apt, arr_apt, etc.).
    """
    if use_stub:
        # ---------- STUB DATA (edit as desired) ----------
        # 6 tails, uneven leg counts, mixed timezones
        raw = [
            {
                "tail": "C-GASL",
                "leg_id": "L1",
                "dep_time": f"{target_date}T06:15:00-04:00",
                "dep_tz": "America/New_York",
                "workflowCustomName": "FEX Guaranteed",
            },
            {"tail": "C-GASL", "leg_id": "L2", "dep_time": f"{target_date}T09:40:00-04:00", "dep_tz": "America/New_York"},

            {"tail": "C-FLYR", "leg_id": "L3", "dep_time": f"{target_date}T05:55:00-07:00", "dep_tz": "America/Los_Angeles"},

            {"tail": "C-JETA", "leg_id": "L4", "dep_time": f"{target_date}T07:20:00-06:00", "dep_tz": "America/Denver"},
            {"tail": "C-JETA", "leg_id": "L5", "dep_time": f"{target_date}T12:10:00-06:00", "dep_tz": "America/Denver"},
            {"tail": "C-JETA", "leg_id": "L6", "dep_time": f"{target_date}T18:25:00-06:00", "dep_tz": "America/Denver"},

            {
                "tail": "C-LEGC",
                "leg_id": "L7",
                "dep_time": f"{target_date}T14:45:00+01:00",
                "dep_tz": "Europe/London",
                "workflowCustomName": "Priority Charter",
            },
            {"tail": "C-LEGC", "leg_id": "L8", "dep_time": f"{target_date}T19:30:00+01:00", "dep_tz": "Europe/London"},

            {"tail": "C-CJ25", "leg_id": "L9", "dep_time": f"{target_date}T06:05:00-05:00", "dep_tz": "America/Chicago"},

            {"tail": "C-HAWK", "leg_id": "L10", "dep_time": f"{target_date}T08:00:00-06:00", "dep_tz": "America/Denver"},
            {"tail": "C-HAWK", "leg_id": "L11", "dep_time": f"{target_date}T16:40:00-06:00", "dep_tz": "America/Denver"},
        ]
        filtered_raw, skipped_subcharter = _filter_out_subcharter_rows(raw)
        if skipped_subcharter:
            st.info(
                "Skipped %d leg%s because the workflow contains 'Subcharter'."
                % (skipped_subcharter, "s" if skipped_subcharter != 1 else "")
            )
        metadata: Dict[str, Any] = {}
        if skipped_subcharter:
            metadata["skipped_subcharter_legs"] = skipped_subcharter
        return pd.DataFrame(filtered_raw), metadata, None

    # ---------- REAL FETCH ----------
    if not fl3xx_settings:
        st.error("FL3XX API secrets are not configured; falling back to stub data is recommended.")
        return pd.DataFrame(), {}, None

    def _coerce_bool(value: Any, default: bool) -> bool:
        if isinstance(value, bool):
            return value
        if isinstance(value, str):
            return value.strip().lower() in {"1", "true", "yes", "on"}
        return default

    def _coerce_int(value: Any, default: int) -> int:
        try:
            return int(value)
        except (TypeError, ValueError):
            return default

    extra_headers = fl3xx_settings.get("extra_headers")
    if isinstance(extra_headers, dict):
        sanitized_headers = {str(k): str(v) for k, v in extra_headers.items()}
    else:
        sanitized_headers = {}

    extra_params = fl3xx_settings.get("extra_params")
    if isinstance(extra_params, dict):
        sanitized_params = {str(k): str(v) for k, v in extra_params.items()}
    else:
        sanitized_params = {}

    config = Fl3xxApiConfig(
        base_url=str(fl3xx_settings.get("base_url") or DEFAULT_FL3XX_BASE_URL),
        api_token=str(fl3xx_settings.get("api_token")) if fl3xx_settings.get("api_token") else None,
        auth_header=str(fl3xx_settings.get("auth_header")) if fl3xx_settings.get("auth_header") else None,
        auth_header_name=str(fl3xx_settings.get("auth_header_name") or "Authorization"),
        api_token_scheme=str(fl3xx_settings.get("api_token_scheme")) if fl3xx_settings.get("api_token_scheme") else None,
        extra_headers=sanitized_headers,
        verify_ssl=_coerce_bool(fl3xx_settings.get("verify_ssl"), True),
        timeout=_coerce_int(fl3xx_settings.get("timeout"), 30),
        extra_params=sanitized_params,
    )

    try:
        flights, metadata = fetch_flights(
            config,
            from_date=target_date,
            to_date=target_date + timedelta(days=2),
        )
    except Exception as exc:
        st.error(f"Error fetching data from FL3XX API: {exc}")
        return pd.DataFrame(), {}, None

    crew_summary: Optional[Dict[str, Any]] = None
    if fetch_crew:
        crew_summary = enrich_flights_with_crew(config, flights)
        metadata = {**metadata, "crew_summary": crew_summary}

    window_start_utc, window_end_utc = _compute_departure_window_bounds(target_date)
    window_meta = {
        "start": _format_utc(window_start_utc),
        "end": _format_utc(window_end_utc),
    }

    normalized_rows, normalization_stats = _normalize_fl3xx_payload({"items": flights})
    original_normalized_count = len(normalized_rows)
    normalized_rows, skipped_subcharter = _filter_out_subcharter_rows(normalized_rows)
    if skipped_subcharter:
        st.info(
            "Skipped %d leg%s because the workflow contains 'Subcharter'."
            % (skipped_subcharter, "s" if skipped_subcharter != 1 else "")
        )
    normalization_stats["skipped_subcharter"] = skipped_subcharter
    rows, window_stats = _filter_rows_by_departure_window(
        normalized_rows, window_start_utc, window_end_utc
    )
    metadata = {
        **metadata,
        "normalization_stats": normalization_stats,
        "departure_window_utc": window_meta,
        "departure_window_counts": window_stats,
    }
    if skipped_subcharter:
        metadata["skipped_subcharter_legs"] = skipped_subcharter
    if not normalized_rows:
        if original_normalized_count and skipped_subcharter == original_normalized_count:
            return pd.DataFrame(), metadata, crew_summary
        st.warning("FL3XX API returned no recognizable legs for the selected date.")
        return pd.DataFrame(), metadata, crew_summary
    if not rows:
        st.warning(
            "No FL3XX legs depart within the UTC window from %s to %s."
            % (window_meta["start"], window_meta["end"])
        )
        return pd.DataFrame(), metadata, crew_summary

    df = pd.DataFrame(rows)
    df, missing_tz_airports, tz_lookup_used = _apply_airport_timezones(df)

    metadata = {
        **metadata,
        "timezone_lookup_used": tz_lookup_used,
    }
    if missing_tz_airports:
        metadata["missing_dep_tz_airports"] = sorted(missing_tz_airports)

    skipped_tail = normalization_stats.get("skipped_missing_tail", 0)
    skipped_time = normalization_stats.get("skipped_missing_dep_time", 0)
    if skipped_tail or skipped_time:
        skipped_total = skipped_tail + skipped_time
        st.warning(
            "Skipped %d leg%s missing required fields (tail missing: %d, departure time missing: %d)."
            % (
                skipped_total,
                "s" if skipped_total != 1 else "",
                skipped_tail,
                skipped_time,
            )
        )

    if missing_tz_airports:
        sample = ", ".join(sorted(missing_tz_airports))
        if len(sample) > 200:
            sample = sample[:197] + "..."
        message = (
            "Added timezone from airport lookup where possible. Update `%s` to cover: %s"
            % (AIRPORT_TZ_FILENAME, sample)
        )
        if tz_lookup_used:
            st.info(message)
        else:
            st.warning(
                "Unable to infer departure timezones automatically because `%s` was not found. "
                "Sample airports without tz: %s"
                % (AIRPORT_TZ_FILENAME, sample)
            )

    return df, metadata, crew_summary


def _extract_first(obj: Dict[str, Any], *keys: str) -> Any:
    for key in keys:
        if key in obj and obj[key] not in (None, ""):
            return obj[key]
    return None


def _normalize_fl3xx_payload(payload: Any) -> Tuple[List[Dict[str, Any]], Dict[str, int]]:
    """Best-effort normalization of FL3XX flights/legs payload to rows with required fields."""

    def _iterable_items(data: Any) -> List[Dict[str, Any]]:
        if isinstance(data, list):
            return data
        if isinstance(data, dict):
            for key in ("data", "items", "flights", "legs"):
                nested = data.get(key)
                if isinstance(nested, list):
                    return nested
        return []

    items = _iterable_items(payload)
    if not items and isinstance(payload, dict):
        items = [payload]
    elif not items and isinstance(payload, list):
        items = payload

    normalized: List[Dict[str, Any]] = []
    stats = {
        "flights_processed": len(items),
        "candidate_legs": 0,
        "legs_normalized": 0,
        "skipped_missing_tail": 0,
        "skipped_missing_dep_time": 0,
    }
    for flight in items:
        legs = []
        if isinstance(flight, dict):
            legs_data = flight.get("legs")
            if isinstance(legs_data, list) and legs_data:
                legs = legs_data
            else:
                legs = [flight]
        elif isinstance(flight, list):
            legs = flight
        else:
            continue

        flight_tail = {}
        if isinstance(flight, dict):
            flight_tail = flight

        for leg in legs:
            if not isinstance(leg, dict):
                continue
            stats["candidate_legs"] += 1
            tail = _extract_first(
                leg,
                "tail",
                "tailNumber",
                "tail_number",
                "aircraft",
                "aircraftRegistration",
                "registrationNumber",
                "registration",
            )
            if not tail and isinstance(flight_tail, dict):
                tail = _extract_first(
                    flight_tail,
                    "tail",
                    "tailNumber",
                    "tail_number",
                    "aircraft",
                    "aircraftRegistration",
                    "registrationNumber",
                    "registration",
                )
            if isinstance(tail, dict):
                tail = _extract_first(
                    tail,
                    "registrationNumber",
                    "registration",
                    "tailNumber",
                    "tail",
                    "name",
                )

            leg_id = _extract_first(
                leg,
                "id",
                "legId",
                "leg_id",
                "uuid",
                "externalId",
                "external_id",
            )

            dep_time = _extract_first(
                leg,
                "departureTimeUtc",
                "departure_time_utc",
                "departureTime",
                "departure_time",
                "offBlockTimeUtc",
                "scheduledTimeUtc",
                "scheduled_departure_utc",
                "blockOffEstUTC",
                "blockOffUtc",
                "scheduledOffBlockUtc",
                "blockOffTimeUtc",
                "blockOffActualUTC",
            )

            dep_tz = _extract_first(
                leg,
                "departureTimezone",
                "departureTimeZone",
                "departure_timezone",
                "departure_tz",
                "blockOffTimeZone",
                "offBlockTimeZone",
            )
            if not dep_tz:
                dep = leg.get("departure") if isinstance(leg.get("departure"), dict) else {}
                if isinstance(dep, dict):
                    dep_tz = _extract_first(dep, "timezone", "timeZone")

            if not tail:
                stats["skipped_missing_tail"] += 1
                continue
            if not dep_time:
                stats["skipped_missing_dep_time"] += 1
                continue

            def _coerce_name(container: Dict[str, Any], *keys: str) -> Optional[str]:
                value = _extract_first(container, *keys)
                if value is None:
                    return None
                value_str = str(value).strip()
                return value_str or None

            pic_name = _coerce_name(
                leg,
                "picName",
                "pic",
                "pic_name",
                "captainName",
                "captain",
            )
            if not pic_name and isinstance(flight_tail, dict):
                pic_name = _coerce_name(
                    flight_tail,
                    "picName",
                    "pic",
                    "pic_name",
                    "captainName",
                    "captain",
                )

            sic_name = _coerce_name(
                leg,
                "sicName",
                "sic",
                "foName",
                "firstOfficer",
            )
            if not sic_name and isinstance(flight_tail, dict):
                sic_name = _coerce_name(
                    flight_tail,
                    "sicName",
                    "sic",
                    "foName",
                    "firstOfficer",
                )

            workflow_custom_name = _extract_first(
                leg,
                "workflowCustomName",
                "workflow_custom_name",
                "workflowName",
                "workflow",
            )
            if not workflow_custom_name and isinstance(flight_tail, dict):
                workflow_custom_name = _extract_first(
                    flight_tail,
                    "workflowCustomName",
                    "workflow_custom_name",
                    "workflowName",
                    "workflow",
                )

            row = {
                "tail": str(tail),
                "leg_id": str(leg_id) if leg_id is not None else str(len(normalized) + 1),
                "dep_time": dep_time,
                "dep_tz": dep_tz,
            }
            if pic_name:
                row["picName"] = pic_name
            if sic_name:
                row["sicName"] = sic_name
            if workflow_custom_name:
                row["workflowCustomName"] = str(workflow_custom_name)

            dep_airport = _extract_first(
                leg,
                "departureAirport",
                "departureAirportCode",
                "departureAirportIcao",
                "departureAirportIata",
                "departureAirportName",
                "departure",
                "airportFrom",
                "fromAirport",
            )
            if isinstance(dep_airport, dict):
                dep_airport = _extract_first(
                    dep_airport,
                    "icao",
                    "iata",
                    "code",
                    "name",
                    "airport",
                )
            if dep_airport:
                row["departure_airport"] = str(dep_airport)

            arr_airport = _extract_first(
                leg,
                "arrivalAirport",
                "arrivalAirportCode",
                "arrivalAirportIcao",
                "arrivalAirportIata",
                "arrivalAirportName",
                "arrival",
                "airportTo",
                "toAirport",
            )
            if isinstance(arr_airport, dict):
                arr_airport = _extract_first(
                    arr_airport,
                    "icao",
                    "iata",
                    "code",
                    "name",
                    "airport",
                )
            if arr_airport:
                row["arrival_airport"] = str(arr_airport)

            if isinstance(leg.get("crewMembers"), list):
                row["crewMembers"] = leg["crewMembers"]
            elif isinstance(flight_tail, dict) and isinstance(flight_tail.get("crewMembers"), list):
                row["crewMembers"] = flight_tail["crewMembers"]

            normalized.append(row)
            stats["legs_normalized"] += 1

    return normalized, stats


def build_tail_packages(df: pd.DataFrame, target_date: date) -> Tuple[List[TailPackage], Set[str]]:
    if df.empty:
        return [], set()
    # Ensure required columns
    required = {"tail", "leg_id", "dep_time"}
    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Missing required columns in data: {missing}")

    df = df.copy()
    df["tail"] = df["tail"].astype(str)

    invalid_tails: Set[str] = set()

    def _valid_tail(value: Any) -> bool:
        tail_str = str(value)
        is_valid = _is_valid_tail_registration(tail_str)
        if not is_valid:
            invalid_tails.add(tail_str.strip())
        return is_valid

    df = df[df["tail"].map(_valid_tail)]
    if df.empty:
        return [], invalid_tails

    # Derive local first departure per tail for the day
    def first_local_for_tail(g: pd.DataFrame) -> datetime:
        # Filter legs that depart on target_date in their *local* timezone
        times_local: List[datetime] = []
        for _, row in g.iterrows():
            dt = _safe_parse_dt(str(row["dep_time"]))
            tz_name = str(row.get("dep_tz", "")) or None
            dt_local = _to_local(dt, tz_name)
            if dt_local.date() == target_date:
                times_local.append(dt_local)
        if not times_local:
            # If none match exactly by local date, fall back to min local
            for _, row in g.iterrows():
                dt = _safe_parse_dt(str(row["dep_time"]))
                tz_name = str(row.get("dep_tz", "")) or None
                times_local.append(_to_local(dt, tz_name))
        return min(times_local)

    airport_lookup = _load_airport_metadata_lookup()

    packages: List[TailPackage] = []
    for tail, g in df.groupby("tail", sort=False):
        # Limit to target_date legs (by local date)
        legs_rows: List[Dict[str, Any]] = []
        all_rows: List[Dict[str, Any]] = []
        priority_values: Set[str] = set()
        for _, row in g.iterrows():
            row_dict = row.to_dict()
            is_customs = _is_customs_leg(row_dict, airport_lookup)
            row_dict["is_customs_leg"] = is_customs
            all_rows.append(row_dict)
            priority_label = _priority_label(row_dict.get("workflowCustomName"))
            if priority_label:
                priority_values.add(priority_label)
            dt = _safe_parse_dt(str(row_dict["dep_time"]))
            tz_name = str(row_dict.get("dep_tz", "")) or None
            dt_local = _to_local(dt, tz_name)
            if dt_local.date() == target_date:
                legs_rows.append(row_dict)
        # If none strictly on target_date by local, treat all as same-day package
        if not legs_rows:
            legs_rows = all_rows
        first_dt = first_local_for_tail(pd.DataFrame(legs_rows))
        customs_count = sum(1 for leg in legs_rows if leg.get("is_customs_leg"))
        workload = 0.0
        for leg in legs_rows:
            workload += _CUSTOMS_WORKLOAD_MULTIPLIER if leg.get("is_customs_leg") else 1.0
        packages.append(
            TailPackage(
                tail=str(tail),
                legs=len(legs_rows),
                workload=workload,
                first_local_dt=first_dt,
                sample_legs=legs_rows[:3],
                has_priority=bool(priority_values),
                priority_labels=sorted(priority_values),
                customs_legs=customs_count,
            )
        )
    return packages, invalid_tails


def assign_round_robin_by_first(packages: List[TailPackage], labels: List[str]) -> Dict[str, List[TailPackage]]:
    packages_sorted = sorted(packages, key=lambda p: p.first_local_dt)
    buckets: Dict[str, List[TailPackage]] = {lab: [] for lab in labels}
    for i, pkg in enumerate(packages_sorted):
        label = labels[i % len(labels)]
        buckets[label].append(pkg)
    return buckets


def assign_balanced_by_legs(packages: List[TailPackage], labels: List[str]) -> Dict[str, List[TailPackage]]:
    # Greedy bin-pack: biggest packages first → assign to bucket with lowest total legs
    buckets: Dict[str, List[TailPackage]] = {lab: [] for lab in labels}
    totals = {lab: 0.0 for lab in labels}

    def _workload(pkg: TailPackage) -> float:
        return pkg.workload if pkg.workload else float(pkg.legs)

    for pkg in sorted(packages, key=lambda p: _workload(p), reverse=True):
        # choose label with smallest total, then smallest count, then order
        label = sorted(
            labels,
            key=lambda lab: (totals[lab], len(buckets[lab]), labels.index(lab)),
        )[0]
        buckets[label].append(pkg)
        totals[label] += _workload(pkg)
    return buckets


def _offset_hours(dt: datetime) -> float:
    offset = dt.utcoffset()
    if offset is None:
        return 0.0
    return offset.total_seconds() / 3600


def assign_preference_weighted(packages: List[TailPackage], labels: List[str]) -> Dict[str, List[TailPackage]]:
    if not packages or not labels:
        return {lab: [] for lab in labels}

    offsets = [_offset_hours(pkg.first_local_dt) for pkg in packages]
    min_off, max_off = min(offsets), max(offsets)
    def _workload(pkg: TailPackage) -> float:
        return pkg.workload if pkg.workload else float(pkg.legs)

    total_workload = sum(_workload(pkg) for pkg in packages)
    avg_workload = total_workload / len(labels) if labels else 0.0
    # Keep a small tolerance so we still respect the east↔west preference, but
    # not at the expense of an even split.
    tolerance = max(0.5, round(avg_workload * 0.25, 2)) if avg_workload else 0.5
    if len(labels) == 1:
        targets = [max_off]
    elif max_off == min_off:
        targets = [max_off for _ in labels]
    else:
        step = (max_off - min_off) / (len(labels) - 1)
        targets = [max_off - step * idx for idx in range(len(labels))]

    buckets: Dict[str, List[TailPackage]] = {lab: [] for lab in labels}
    totals = {lab: 0.0 for lab in labels}

    span = max_off - min_off
    center = min_off + span / 2 if span else min_off

    for pkg in sorted(packages, key=lambda p: p.first_local_dt):
        pkg_offset = _offset_hours(pkg.first_local_dt)
        min_total = min(totals.values())
        eligible_labels = [lab for lab in labels if totals[lab] <= min_total + tolerance]
        if not eligible_labels:
            eligible_labels = labels

        def score(lab: str) -> tuple[float, float, float, int, int]:
            target = targets[labels.index(lab)]
            if span:
                half_span = span / 2 or 1
                # Increase the penalty for extreme east/west packages so that
                # far-west departures lean harder toward later shifts (and far-east
                # toward earlier ones) than mid-range timezones.
                normalized_extremity = min(abs(pkg_offset - center) / half_span, 2)
                weight = 1 + normalized_extremity
            else:
                weight = 1.0
            tz_penalty = abs(pkg_offset - target) * weight
            projected_total = totals[lab] + _workload(pkg)
            return (
                round(abs(projected_total - avg_workload), 4),
                round(projected_total - min_total, 4),
                round(tz_penalty, 4),
                len(buckets[lab]),
                labels.index(lab),
            )

        label = min(eligible_labels, key=score)
        buckets[label].append(pkg)
        totals[label] += _workload(pkg)

    return buckets


def buckets_to_df(buckets: Dict[str, List[TailPackage]]) -> pd.DataFrame:
    rows = []
    for label, pkgs in buckets.items():
        for pkg in sorted(pkgs, key=lambda p: (p.first_local_dt, p.tail)):
            rows.append({
                "Shift": label,
                "Tail": pkg.tail,
                "Legs": pkg.legs,
                "Customs Legs": pkg.customs_legs,
                "Workload": round(pkg.workload, 2),
                "First Local Dep": pkg.first_local_dt.strftime("%Y-%m-%d %H:%M %Z"),
                "Priority Flight": "Yes" if pkg.has_priority else "No",
                "Priority Detail": ", ".join(pkg.priority_labels) if pkg.priority_labels else "",
            })
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.sort_values(["Shift", "First Local Dep", "Tail"]).reset_index(drop=True)
    return df


def summarize(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame()
    agg = (
        df.groupby("Shift")
        .agg(
            Tails=("Tail", "count"),
            Legs=("Legs", "sum"),
            Customs=("Customs Legs", "sum"),
            Workload=("Workload", "sum"),
        )
        .reset_index()
    ).rename(columns={"Customs": "Customs Legs"})
    agg["Workload"] = agg["Workload"].round(2)
    # Add spread metrics
    total_workload = agg["Workload"].sum()
    total_shifts = agg.shape[0]
    target = total_workload / total_shifts if total_shifts else 0
    agg["Δ Workload vs Even"] = (agg["Workload"] - target).round(2)
    return agg


_DOCX_HEADERS = [
    "TAIL #",
    "CREW PIC",
    "CREW SIC",
    "FUEL",
    "CUSTOMS",
    "SLOT / PPR",
    "FLIGHT PLANS",
    "CREW BRIEF",
    "CONFIRMATION PIC",
    "CONFIRMATION SIC",
    "CHECK LIST",
    "RELEASE",
    "NOTES",
    "Priority Status",
]

_CHECKMARK = "✓"


def _apply_landscape(document: Document) -> None:
    for section in document.sections:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width


def _initialize_briefing_document(target_date: date) -> Document:
    document = Document()
    document.core_properties.title = f"{target_date} Shift Briefing"
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(8)

    title_para = document.add_paragraph(f"Daily Flight Sheet – {target_date}")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_para.runs:
        title_run = title_para.runs[0]
    else:
        title_run = title_para.add_run(f"Daily Flight Sheet – {target_date}")
    title_run.font.size = Pt(16)
    title_run.bold = True

    _apply_landscape(document)
    return document


def _add_shift_table(
    document: Document,
    label: str,
    pkgs: List[TailPackage],
    priority_details: Dict[str, str],
) -> None:
    sorted_pkgs = sorted(pkgs, key=lambda p: (p.first_local_dt, p.tail))
    table_rows = len(sorted_pkgs) + 3  # header row + column headers + data + footer
    table = document.add_table(rows=table_rows, cols=len(_DOCX_HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Shift label header row spanning all columns
    top_cell = table.rows[0].cells[0]
    for merge_idx in range(1, len(_DOCX_HEADERS)):
        top_cell = top_cell.merge(table.rows[0].cells[merge_idx])
    top_paragraph = top_cell.paragraphs[0]
    top_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = top_paragraph.add_run(label)
    run.bold = True
    run.font.size = Pt(14)

    # Column headers
    header_row = table.rows[1]
    for col_idx, header_text in enumerate(_DOCX_HEADERS):
        header_cell = header_row.cells[col_idx]
        header_cell.text = header_text
        header_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        header_paragraph = header_cell.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_paragraph.runs[0].font.bold = True

    # Data rows
    for row_offset, pkg in enumerate(sorted_pkgs):
        row = table.rows[row_offset + 2]
        pic_name, sic_name = _crew_names_from_package(pkg)
        values = [""] * len(_DOCX_HEADERS)
        values[0] = pkg.tail
        values[1] = pic_name
        values[2] = sic_name
        detail = priority_details.get(pkg.tail, "")
        cleaned_detail = ""
        if detail and not detail.lower().startswith("priority"):
            cleaned_detail = detail
        elif detail:
            cleaned_detail = detail.replace("priority", "", 1).strip() or detail
        if pkg.has_priority:
            values[13] = _CHECKMARK
            if cleaned_detail:
                values[13] = f"{values[13]} {cleaned_detail}".strip()
        elif cleaned_detail:
            values[13] = cleaned_detail
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if col_idx in {0, 13}:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Footer row for positioning/notes
    footer_row = table.rows[-1]
    positioning_cell = footer_row.cells[0]
    for merge_idx in range(1, max(1, len(_DOCX_HEADERS) // 2)):
        positioning_cell = positioning_cell.merge(footer_row.cells[merge_idx])
    positioning_cell.text = "POSITIONING:"
    positioning_cell.paragraphs[0].runs[0].bold = True

    notes_start = len(_DOCX_HEADERS) // 2
    notes_cell = footer_row.cells[notes_start]
    for merge_idx in range(notes_start + 1, len(_DOCX_HEADERS)):
        notes_cell = notes_cell.merge(footer_row.cells[merge_idx])
    notes_cell.text = "ADDITIONAL NOTES:"
    notes_cell.paragraphs[0].runs[0].bold = True


def _document_to_bytes(document: Document) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _label_slug(label: str) -> str:
    slug = re.sub(r"[^A-Za-z0-9]+", "_", label.strip()).strip("_")
    return slug or "shift"


def build_shift_briefing_docs(
    target_date: date,
    labels: List[str],
    buckets: Dict[str, List[TailPackage]],
    priority_details: Dict[str, str],
) -> tuple[bytes, Dict[str, bytes]]:
    combined_document = _initialize_briefing_document(target_date)

    for idx, label in enumerate(labels):
        pkgs = buckets.get(label, [])
        if idx > 0:
            combined_document.add_paragraph("")
        _add_shift_table(combined_document, label, pkgs, priority_details)

    combined_payload = _document_to_bytes(combined_document)

    per_shift_payloads: Dict[str, bytes] = {}
    for label in labels:
        shift_document = _initialize_briefing_document(target_date)
        _add_shift_table(shift_document, label, buckets.get(label, []), priority_details)
        per_shift_payloads[label] = _document_to_bytes(shift_document)

    return combined_payload, per_shift_payloads


# ----------------------------
# Sidebar: Inputs
# ----------------------------
st.sidebar.header("Inputs")

fl3xx_cfg: Dict[str, Any] = {}
try:
    if "fl3xx_api" in st.secrets:
        cfg = st.secrets["fl3xx_api"]
        if isinstance(cfg, Mapping):
            fl3xx_cfg = {str(k): cfg[k] for k in cfg}
        elif isinstance(cfg, dict):
            fl3xx_cfg = dict(cfg)
except Exception:
    # Accessing secrets outside Streamlit Cloud may raise; ignore gracefully.
    fl3xx_cfg = {}

has_live_credentials = bool(fl3xx_cfg.get("api_token") or fl3xx_cfg.get("auth_header"))

use_stub = st.sidebar.toggle(
    "Use stub data",
    value=not has_live_credentials,
    help="Uncheck to use your real FL3XX API (credentials stored in Streamlit secrets).",
    disabled=not has_live_credentials,
)

if not has_live_credentials:
    st.sidebar.info(
        "Add your FL3XX credentials to `.streamlit/secrets.toml` under `[fl3xx_api]` to enable live fetching.",
    )
else:
    st.sidebar.success("Using FL3XX credentials from Streamlit secrets.")

fetch_crew_default = bool(fl3xx_cfg.get("fetch_crew", True))
fetch_crew = st.sidebar.toggle(
    "Fetch crew details",
    value=fetch_crew_default,
    help="Retrieve crew information (PIC/SIC) for each flight. Requires additional API calls.",
    disabled=use_stub,
)

num_people = st.sidebar.number_input("Number of on-duty people", min_value=1, max_value=12, value=4, step=1)

default_labels = _default_shift_labels(int(num_people))
labels = []
for i in range(int(num_people)):
    lbl = st.sidebar.text_input(
        f"Label for person {i+1}",
        value=default_labels[i] if i < len(default_labels) else f"Shift {i+1}",
    )
    labels.append(lbl or f"Shift {i+1}")

# Date selection (default = two days ahead in local Mountain time)
selected_date = st.sidebar.date_input("Target date", value=_default_target_date())


# ----------------------------
# Main Action
# ----------------------------
fetch_col, reset_col = st.columns([4, 1])
with fetch_col:
    if st.button("🔄 Fetch & Assign", use_container_width=True):
        st.session_state["_run"] = True
with reset_col:
    if st.button("🧹 Clear cache", use_container_width=True):
        fetch_next_day_legs.clear()
        st.session_state.pop("_run", None)
        st.session_state["_cache_cleared"] = True

if st.session_state.pop("_cache_cleared", False):
    st.success("Cached data cleared. Fetch again to pull fresh data.")

# ----------------------------
# Processing & Output
# ----------------------------
if st.session_state.get("_run"):
    legs_df, _, crew_summary = fetch_next_day_legs(
        selected_date,
        use_stub=use_stub,
        fl3xx_settings=fl3xx_cfg if not use_stub else None,
        fetch_crew=bool(fetch_crew and not use_stub),
    )

    if legs_df.empty:
        st.warning("No legs returned for the selected date.")
        st.stop()

    if crew_summary and crew_summary.get("fetched"):
        st.sidebar.metric("Crew lookups", int(crew_summary["fetched"]))
        if crew_summary.get("errors"):
            st.sidebar.warning(f"Crew errors: {len(crew_summary['errors'])}")

    packages, invalid_tails = build_tail_packages(legs_df, selected_date)

    if invalid_tails:
        ignored = sorted(t for t in invalid_tails if t)
        if ignored:
            preview = ", ".join(ignored[:6])
            if len(ignored) > 6:
                preview += ", ..."
            st.info(
                "Ignored %d tail%s without an official registration: %s"
                % (
                    len(ignored),
                    "s" if len(ignored) != 1 else "",
                    preview,
                )
            )

    if not packages:
        st.info("No tail packages found for the selected date.")
        st.stop()

    priority_packages = [pkg for pkg in packages if pkg.has_priority]
    priority_tails = [pkg.tail for pkg in priority_packages]
    priority_details = {
        pkg.tail: ", ".join(pkg.priority_labels) if pkg.priority_labels else ""
        for pkg in priority_packages
    }

    st.subheader("Assignments")

    buckets = assign_preference_weighted(packages, labels)

    # Display per-shift tables
    tabs = st.tabs(labels)
    for i, lab in enumerate(labels):
        with tabs[i]:
            pkgs = buckets.get(lab, [])
            df = buckets_to_df({lab: pkgs})
            if df.empty:
                st.write("No tails assigned.")
            else:
                st.dataframe(df, use_container_width=True, hide_index=True)
                total_legs = int(df["Legs"].sum())
                total_workload = round(float(df["Workload"].sum()), 2)
                customs_legs = int(df["Customs Legs"].sum())
                total_tails = int(df.shape[0])
                priority_total = int(sum(1 for p in pkgs if p.has_priority))
                col1, col2, col3, col4, col5 = st.columns(5)
                with col1:
                    st.metric("Total legs", total_legs)
                with col2:
                    st.metric("Workload-adjusted legs", total_workload)
                with col3:
                    st.metric("Customs legs", customs_legs)
                with col4:
                    st.metric("Tails", total_tails)
                with col5:
                    st.metric("Priority tails", priority_total)

    # Combined view
    combined_df = buckets_to_df(buckets)

    if priority_tails:
        detail_list = [
            f"{tail} ({priority_details[tail]})" if priority_details[tail] else tail
            for tail in priority_tails
        ]
        st.warning(
            "Priority flights detected for: " + ", ".join(detail_list)
        )

    # Summary
    st.subheader("Summary")
    summary_df = summarize(combined_df)
    st.dataframe(summary_df, use_container_width=True, hide_index=True)

    # Downloads
    doc_payload, per_shift_docs = build_shift_briefing_docs(
        selected_date, labels, buckets, priority_details
    )
    st.download_button(
        label="⬇️ Download daily flight sheet (DOCX)",
        data=doc_payload,
        file_name=f"daily_flight_sheet_{selected_date}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    if per_shift_docs:
        st.markdown("#### Individual shift documents")
        columns = st.columns(min(3, len(per_shift_docs)) or 1)
        for idx, label in enumerate(labels):
            payload = per_shift_docs.get(label)
            if not payload:
                continue
            column = columns[idx % len(columns)]
            with column:
                st.download_button(
                    label=f"⬇️ {label}",
                    data=payload,
                    file_name=f"daily_flight_sheet_{selected_date}_{_label_slug(label)}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
    st.download_button(
        label="⬇️ Download assignments (CSV)",
        data=combined_df.to_csv(index=False).encode("utf-8"),
        file_name=f"tail_assignments_{selected_date}.csv",
        mime="text/csv",
        use_container_width=True,
    )

    st.success("Done. Adjust inputs and re-run as needed.")

